"""
THEeye - AI-Assisted Research Platform
FastAPI Application Server with Authentication

Features:
  - User registration and login (JWT-style tokens)
  - Admin backend with user management, config, and content editing
  - Literature search across OpenAlex, Crossref, Semantic Scholar,
    Google Scholar, EconPapers/RePEc, ERIC
  - Q1/Q2/Q3 quartile filtering
  - Structured data extraction
  - Academic draft generation with citations
  - Quick-generate: single prompt -> final output
  - Data analysis: R, RStudio, Stata 18 code generation
  - Writing tools: Grammarly, QuillBot, Paperpal integration
  - Reference management: Mendeley, Zotero, NotePal export
  - Provenance tracking and integrity audit
"""

import os
import re
import tempfile
from pathlib import Path
from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException, Request, Depends, Header
from fastapi.responses import HTMLResponse, JSONResponse, PlainTextResponse, Response
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware

load_dotenv()

from .models import (
    SearchRequest, SearchResponse,
    ExtractionRequest, ExtractionBatchRequest, ExtractedData,
    DraftRequest, DraftResponse,
    RegisterRequest, LoginRequest, AuthResponse,
    QuickGenerateRequest, QuickGenerateResponse,
    AnalysisRequest, AnalysisResponse, AnalysisRecommendationRequest,
    WritingAnalysisRequest,
    ReferenceExportRequest, ReferenceExportResponse, CitationFormatRequest,
    CitationVerificationRequest,
    ConfigUpdateRequest, ContentUpdateRequest, UserRoleUpdateRequest,
    PasswordChangeRequest, AdminResetPasswordRequest, ProfileUpdateRequest,
    AiTestRequest,
)
from .services import unified_search
from .extraction import extract_from_paper, extract_batch, build_comparison_table
from .drafting import generate_draft
from .audit import create_session, get_session, list_sessions
from .quartiles import get_all_journals, lookup_quartile
from .auth import (
    register_user, login_user, verify_token, logout_user,
    is_admin, get_all_users, get_user_by_id, update_user_role,
    toggle_user_active, delete_user, update_user_profile,
    change_password, admin_reset_password,
)
from .admin import (
    get_config, update_config, get_feature_flags,
    get_content, update_content, list_content_keys,
    get_stats, record_stat, reset_stats,
    get_database_sources, toggle_database_source,
    get_tool_integrations, toggle_tool_integration,
)
from .analysis import (
    generate_r_code, generate_stata_code, recommend_analysis,
    get_available_methods, ANALYSIS_METHODS,
)
from .writing_tools import (
    get_writing_tools, get_writing_tool,
    analyze_writing, enhance_for_journal,
    fix_grammar, paraphrase_text, enhance_academic, enhance_text_all,
    humanize_text,
)

from .data_analysis import (
    parse_dataset, execute_analysis, extract_online_data,
    suggest_econometric_model, generate_tool_code,
)

from .references import (
    get_reference_managers, get_reference_manager,
    get_export_formats, export_references, format_citations,
    verify_citations,
)
from .models import DocumentExportRequest

from .document_drafting import (
    generate_roadmap, get_document_types, get_research_fields,
    suggest_research_topics, find_open_access_articles, download_pdf,
    OA_SOURCES,
)
from .models import (
    RoadmapRequest, RoadmapResponse,
    TopicSuggestionRequest, TopicSuggestionResponse,
    OpenAccessArticleRequest, OpenAccessArticleResponse,
    DownloadPdfRequest,
)

from .ai_router import (
    generate_text, enhance_text_with_ai, generate_academic_text,
    get_available_providers, get_routing_info,
    get_ai_status, test_provider_connection,
)


# ---------------------------------------------------------------------------
# Document Export Helpers
# ---------------------------------------------------------------------------

import io as _io
import re as _re
import html as _htmlmod

DOCUMENT_FORMATS = {
    "docx": {"name": "Word Document", "extension": ".docx", "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"},
    "pdf": {"name": "PDF Document", "extension": ".pdf", "mime": "application/pdf"},
    "html": {"name": "HTML Document", "extension": ".html", "mime": "text/html"},
    "md": {"name": "Markdown", "extension": ".md", "mime": "text/markdown"},
    "txt": {"name": "Plain Text", "extension": ".txt", "mime": "text/plain"},
}


def _md_to_blocks(text):
    """Parse markdown ## and ### headings into blocks."""
    blocks = []
    for raw in _re.split(r'\n\s*\n', (text or "").strip()):
        block = raw.strip()
        if not block:
            continue
        if block.startswith("### "):
            blocks.append({"type": "heading3", "text": block[4:].strip()})
        elif block.startswith("## "):
            blocks.append({"type": "heading2", "text": block[3:].strip()})
        else:
            blocks.append({"type": "paragraph", "text": " ".join(l.strip() for l in block.splitlines())})
    return blocks


def _meta_line(topic, section_type, word_count, total_sources):
    parts = []
    if topic: parts.append("Topic: " + topic)
    if section_type: parts.append("Section: " + section_type.replace("_", " ").title())
    if word_count: parts.append("Words: " + str(word_count))
    if total_sources: parts.append("Sources: " + str(total_sources))
    return " | ".join(parts) if parts else ""


def _export_docx(text, title, citations, disclaimer, section_type, topic, word_count, total_sources):
    from docx import Document as _Doc
    from docx.shared import Pt as _Pt, Inches as _In, RGBColor as _RGB
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _AL
    doc = _Doc()
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = _Pt(12)
    st.paragraph_format.line_spacing = 1.5
    st.paragraph_format.space_after = _Pt(6)
    tp = doc.add_paragraph()
    tp.alignment = _AL.CENTER
    tr = tp.add_run(title)
    tr.bold = True
    tr.font.size = _Pt(16)
    tr.font.name = "Times New Roman"
    ml = _meta_line(topic, section_type, word_count, total_sources)
    if ml:
        mp = doc.add_paragraph()
        mp.alignment = _AL.CENTER
        mr = mp.add_run(ml)
        mr.italic = True
        mr.font.size = _Pt(10)
        mr.font.color.rgb = _RGB(0x71, 0x80, 0x96)
    doc.add_paragraph()
    for block in _md_to_blocks(text):
        if block["type"] == "heading2":
            h = doc.add_heading(block["text"], level=2)
            for r in h.runs: r.font.name = "Times New Roman"
        elif block["type"] == "heading3":
            h = doc.add_heading(block["text"], level=3)
            for r in h.runs: r.font.name = "Times New Roman"
        else:
            doc.add_paragraph(block["text"])
    if citations:
        doc.add_paragraph()
        rh = doc.add_heading("References", level=2)
        for r in rh.runs: r.font.name = "Times New Roman"
        for i, c in enumerate(citations, 1):
            authors = c.get("authors", "Unknown") if isinstance(c, dict) else str(c)
            year = c.get("year", "n.d.") if isinstance(c, dict) else ""
            ctitle = c.get("title", "") if isinstance(c, dict) else ""
            journal = c.get("journal", "") if isinstance(c, dict) else ""
            doi = c.get("doi", "") if isinstance(c, dict) else ""
            ref = "[{}] {} ({}). {}. {}.".format(i, authors, year, ctitle, journal)
            if doi: ref += " https://doi.org/" + doi
            p = doc.add_paragraph(ref)
            p.paragraph_format.space_after = _Pt(4)
            p.paragraph_format.left_indent = _In(0.5)
            p.paragraph_format.first_line_indent = _In(-0.5)
    if disclaimer:
        doc.add_paragraph()
        dp = doc.add_paragraph()
        dr = dp.add_run("Integrity Notice: " + disclaimer)
        dr.italic = True
        dr.font.size = _Pt(9)
        dr.font.color.rgb = _RGB(0x74, 0x42, 0x10)
    buf = _io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _export_pdf(text, title, citations, disclaimer, section_type, topic, word_count, total_sources):
    from reportlab.lib.pagesizes import A4 as _A4
    from reportlab.lib.styles import getSampleStyleSheet as _gss, ParagraphStyle as _PS
    from reportlab.lib.units import inch as _inch
    from reportlab.lib.colors import HexColor as _HC
    from reportlab.platypus import SimpleDocTemplate as _SDT, Paragraph as _P, Spacer as _S
    from reportlab.lib.enums import TA_CENTER as _TC, TA_JUSTIFY as _TJ
    buf = _io.BytesIO()
    doc = _SDT(buf, pagesize=_A4, leftMargin=_inch, rightMargin=_inch, topMargin=_inch, bottomMargin=_inch)
    ss = _gss()
    def _esc(s): return _htmlmod.escape(s or "")
    styles = {
        "title": _PS("T", parent=ss["Title"], fontSize=18, alignment=_TC, spaceAfter=6),
        "meta": _PS("M", parent=ss["Normal"], fontSize=9, textColor=_HC("#718096"), alignment=_TC, spaceAfter=14),
        "h2": _PS("H2", parent=ss["Heading2"], fontSize=14, spaceBefore=12, spaceAfter=6),
        "h3": _PS("H3", parent=ss["Heading3"], fontSize=12, spaceBefore=10, spaceAfter=4),
        "body": _PS("B", parent=ss["Normal"], fontSize=11, leading=18, alignment=_TJ, spaceAfter=8),
        "ref": _PS("R", parent=ss["Normal"], fontSize=9, leading=13, leftIndent=18, firstLineIndent=-18, spaceAfter=4),
        "disc": _PS("D", parent=ss["Normal"], fontSize=8, textColor=_HC("#744210"), spaceBefore=12),
    }
    story = [_P(_esc(title), styles["title"])]
    ml = _meta_line(topic, section_type, word_count, total_sources)
    if ml: story.append(_P(_esc(ml), styles["meta"]))
    for block in _md_to_blocks(text):
        et = _esc(block["text"])
        if block["type"] == "heading2": story.append(_P(et, styles["h2"]))
        elif block["type"] == "heading3": story.append(_P(et, styles["h3"]))
        else: story.append(_P(et, styles["body"]))
    if citations:
        story.append(_S(1, 6))
        story.append(_P("References", styles["h2"]))
        for i, c in enumerate(citations, 1):
            authors = c.get("authors", "Unknown") if isinstance(c, dict) else str(c)
            year = c.get("year", "n.d.") if isinstance(c, dict) else ""
            ctitle = c.get("title", "") if isinstance(c, dict) else ""
            journal = c.get("journal", "") if isinstance(c, dict) else ""
            doi = c.get("doi", "") if isinstance(c, dict) else ""
            ref = "[{}] {} ({}). {}. {}.".format(i, _esc(authors), year, _esc(ctitle), _esc(journal))
            if doi: ref += ' <link href="https://doi.org/{}">https://doi.org/{}</link>'.format(doi, doi)
            story.append(_P(ref, styles["ref"]))
    if disclaimer:
        story.append(_S(1, 6))
        story.append(_P("<b>Integrity Notice:</b> " + _esc(disclaimer), styles["disc"]))
    doc.build(story)
    return buf.getvalue()


def _export_html(text, title, citations, disclaimer, section_type, topic, word_count, total_sources):
    blocks = _md_to_blocks(text)
    parts = []
    for b in blocks:
        t = _htmlmod.escape(b["text"])
        if b["type"] == "heading2": parts.append("<h2>{}</h2>".format(t))
        elif b["type"] == "heading3": parts.append("<h3>{}</h3>".format(t))
        else: parts.append("<p>{}</p>".format(t))
    body = "\n".join(parts)
    refs = ""
    if citations:
        items = []
        for i, c in enumerate(citations, 1):
            authors = _htmlmod.escape(c.get("authors", "Unknown") if isinstance(c, dict) else str(c))
            year = c.get("year", "n.d.") if isinstance(c, dict) else ""
            ctitle = _htmlmod.escape(c.get("title", "") if isinstance(c, dict) else "")
            journal = _htmlmod.escape(c.get("journal", "") if isinstance(c, dict) else "")
            doi = c.get("doi", "") if isinstance(c, dict) else ""
            link = ' <a href="https://doi.org/{}">https://doi.org/{}</a>'.format(doi, doi) if doi else ""
            items.append("<li>{} ({}). {}. <em>{}</em>.{}</li>".format(authors, year, ctitle, journal, link))
        refs = '<h2>References</h2><ol class="refs">\n' + "\n".join(items) + "\n</ol>"
    ml = _meta_line(topic, section_type, word_count, total_sources)
    meta_html = '<p class="meta">{}</p>'.format(_htmlmod.escape(ml)) if ml else ""
    disc_html = '<div class="disc"><strong>Integrity Notice:</strong> {}</div>'.format(_htmlmod.escape(disclaimer)) if disclaimer else ""
    etitle = _htmlmod.escape(title)
    doc = '<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8"><title>' + etitle + '</title><style>body{font-family:Times New Roman,serif;max-width:800px;margin:2rem auto;padding:2rem;line-height:1.7;color:#2d3748;}h1{font-size:1.6rem;text-align:center;color:#1a365d;}.meta{text-align:center;font-size:.85rem;color:#718096;font-style:italic;margin-bottom:2rem;}h2{font-size:1.3rem;color:#1a365d;margin-top:1.5rem;}h3{font-size:1.1rem;color:#2c5282;margin-top:1rem;}p{text-align:justify;}.refs{padding-left:1.5rem;font-size:.9rem;}.refs li{margin-bottom:.5rem;}.refs a{color:#3182ce;}.disc{margin-top:2rem;padding:1rem;background:#fffaf0;border-left:4px solid #d69e2e;font-size:.85rem;color:#744210;}</style></head><body><h1>' + etitle + '</h1>' + meta_html + body + refs + disc_html + '</body></html>'
    return doc.encode("utf-8")


def _export_md(text, title, citations, disclaimer, section_type, topic, word_count, total_sources):
    lines = ["# " + title, ""]
    ml = _meta_line(topic, section_type, word_count, total_sources)
    if ml: lines += ["*" + ml + "*", ""]
    lines += [text, ""]
    if citations:
        lines += ["---", "", "## References", ""]
        for i, c in enumerate(citations, 1):
            authors = c.get("authors", "Unknown") if isinstance(c, dict) else str(c)
            year = c.get("year", "n.d.") if isinstance(c, dict) else ""
            ctitle = c.get("title", "") if isinstance(c, dict) else ""
            journal = c.get("journal", "") if isinstance(c, dict) else ""
            doi = c.get("doi", "") if isinstance(c, dict) else ""
            ref = "{}. {} ({}). {}. *{}*.".format(i, authors, year, ctitle, journal)
            if doi: ref += " [https://doi.org/{}](https://doi.org/{})".format(doi, doi)
            lines.append(ref)
        lines.append("")
    if disclaimer: lines += ["---", "", "> **Integrity Notice:** " + disclaimer, ""]
    return "\n".join(lines).encode("utf-8")


def _export_txt(text, title, citations, disclaimer, section_type, topic, word_count, total_sources):
    lines = [title, "=" * max(len(title), 3), ""]
    ml = _meta_line(topic, section_type, word_count, total_sources)
    if ml: lines += [ml, ""]
    lines += [text, ""]
    if citations:
        lines += ["References", "-" * 10]
        for i, c in enumerate(citations, 1):
            authors = c.get("authors", "Unknown") if isinstance(c, dict) else str(c)
            year = c.get("year", "n.d.") if isinstance(c, dict) else ""
            ctitle = c.get("title", "") if isinstance(c, dict) else ""
            journal = c.get("journal", "") if isinstance(c, dict) else ""
            doi = c.get("doi", "") if isinstance(c, dict) else ""
            ref = "[{}] {} ({}). {}. {}.".format(i, authors, year, ctitle, journal)
            if doi: ref += " https://doi.org/" + doi
            lines.append(ref)
        lines.append("")
    if disclaimer: lines += ["-" * 40, "Integrity Notice: " + disclaimer, ""]
    return "\n".join(lines).encode("utf-8")


def _export_document(text, fmt, title, citations, disclaimer, section_type, topic, word_count, total_sources):
    fmt = fmt.lower().strip()
    if fmt not in DOCUMENT_FORMATS:
        raise ValueError("Unsupported format: '{}'. Supported: {}".format(fmt, list(DOCUMENT_FORMATS.keys())))
    safe_topic = _re.sub(r'[^a-zA-Z0-9]', '_', topic or "research")[:50].strip("_") or "research"
    ext = DOCUMENT_FORMATS[fmt]["extension"]
    filename = "THEeye_" + safe_topic + ext
    mime = DOCUMENT_FORMATS[fmt]["mime"]
    kwargs = dict(text=text, title=title, citations=citations, disclaimer=disclaimer, section_type=section_type, topic=topic, word_count=word_count, total_sources=total_sources)
    if fmt == "docx": return _export_docx(**kwargs), filename, mime
    elif fmt == "pdf": return _export_pdf(**kwargs), filename, mime
    elif fmt == "html": return _export_html(**kwargs), filename, mime
    elif fmt == "md": return _export_md(**kwargs), filename, mime
    elif fmt == "txt": return _export_txt(**kwargs), filename, mime


# ---------------------------------------------------------------------------
# App setup
# ---------------------------------------------------------------------------

app = FastAPI(
    title="THEeye - AI-Assisted Research Platform",
    description="Literature discovery, data extraction, and drafting with full integrity compliance.",
    version="2.0.0",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

FRONTEND_DIR = Path(__file__).parent.parent / "frontend"
if FRONTEND_DIR.exists():
    app.mount("/static", StaticFiles(directory=str(FRONTEND_DIR)), name="static")


# ---------------------------------------------------------------------------
# Auth dependency
# ---------------------------------------------------------------------------

async def get_current_user(authorization: str | None = Header(None)):
    """Extract and verify the bearer token from the Authorization header."""
    if not authorization:
        return None
    if authorization.startswith("Bearer "):
        token = authorization[7:]
    else:
        token = authorization
    return verify_token(token)


async def require_auth(authorization: str | None = Header(None)):
    """Require authentication — raises 401 if not authenticated."""
    user = await get_current_user(authorization)
    if not user:
        raise HTTPException(status_code=401, detail="Authentication required. Please login.")
    if not user.is_active:
        raise HTTPException(status_code=403, detail="Account deactivated. Contact administrator.")
    return user


async def require_admin(user=Depends(require_auth)):
    """Require admin role — raises 403 if not admin."""
    if not is_admin(user):
        raise HTTPException(status_code=403, detail="Admin access required.")
    return user


async def _require_auth(request: Request):
    """Manual auth check for Request-based endpoints (non-Depends pattern).
    Returns a JSONResponse error if not authenticated, or None if OK."""
    auth = request.headers.get("authorization")
    if not auth:
        return JSONResponse(status_code=401, content={"detail": "Authentication required. Please login."})
    token = auth[7:] if auth.startswith("Bearer ") else auth
    user = verify_token(token)
    if not user:
        return JSONResponse(status_code=401, content={"detail": "Invalid or expired token."})
    if not user.is_active:
        return JSONResponse(status_code=403, content={"detail": "Account deactivated."})
    return None


# ---------------------------------------------------------------------------
# Frontend
# ---------------------------------------------------------------------------

@app.get("/", response_class=HTMLResponse)
async def index():
    index_path = FRONTEND_DIR / "index.html"
    if index_path.exists():
        return HTMLResponse(content=index_path.read_text(encoding="utf-8"))
    return HTMLResponse(content="<h1>THEeye</h1><p>Frontend not found.</p>")


@app.get("/THEeye", response_class=HTMLResponse)
async def index_theeye():
    """Serve the app at /THEeye path."""
    index_path = FRONTEND_DIR / "index.html"
    if index_path.exists():
        return HTMLResponse(content=index_path.read_text(encoding="utf-8"))
    return HTMLResponse(content="<h1>THEeye</h1><p>Frontend not found.</p>")


# ---------------------------------------------------------------------------
# Health & Info
# ---------------------------------------------------------------------------

@app.get("/api/health")
async def health():
    return {"status": "ok", "platform": "THEeye", "version": "3.0.0"}


@app.get("/api/journals")
async def journals():
    return {"journals": get_all_journals(), "total": len(get_all_journals())}


# ---------------------------------------------------------------------------
# Authentication Routes
# ---------------------------------------------------------------------------

@app.post("/api/auth/register", response_model=AuthResponse)
async def register(request: RegisterRequest):
    """Register a new user account."""
    # Check if public registration is allowed
    config = get_config()
    if not config.get("allow_public_registration", True):
        raise HTTPException(status_code=403, detail="Public registration is disabled. Contact administrator.")
    try:
        result = register_user(
            email=request.email,
            password=request.password,
            name=request.name,
            institution=request.institution,
            research_field=request.research_field,
        )
        record_stat("user_registered")
        return result
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.post("/api/auth/login", response_model=AuthResponse)
async def login(request: LoginRequest):
    """Login with email and password."""
    try:
        result = login_user(request.email, request.password, remember=request.remember)
        return result
    except ValueError as e:
        raise HTTPException(status_code=401, detail=str(e))


@app.post("/api/auth/logout")
async def logout(authorization: str | None = Header(None)):
    """Logout and invalidate the current session token."""
    user = await get_current_user(authorization)
    if not user:
        raise HTTPException(status_code=401, detail="Authentication required.")
    token = authorization[7:] if authorization and authorization.startswith("Bearer ") else authorization
    logout_user(token)
    return {"status": "logged_out"}


@app.get("/api/auth/me")
async def me(user=Depends(require_auth)):
    """Get the current authenticated user's profile."""
    return {
        "id": user.id,
        "email": user.email,
        "name": user.name,
        "institution": user.institution,
        "research_field": user.research_field,
    }


# ---------------------------------------------------------------------------
# Literature Discovery (requires auth)
# ---------------------------------------------------------------------------

@app.post("/api/search", response_model=SearchResponse)
async def search(request: SearchRequest, user=Depends(require_auth)):
    """Search across multiple academic databases with quartile filtering."""
    try:
        record_stat("search", {"databases": request.databases})
        response = await unified_search(request)
        return response
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Search error: {str(e)}")


# ---------------------------------------------------------------------------
# Data Extraction (requires auth)
# ---------------------------------------------------------------------------

@app.post("/api/extract")
async def extract_single(request: ExtractionRequest, user=Depends(require_auth)):
    """Extract structured data from a single paper."""
    try:
        return extract_from_paper(request.paper, use_llm=request.use_llm)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Extraction error: {str(e)}")


@app.post("/api/extract/batch")
async def extract_multiple(request: ExtractionBatchRequest, user=Depends(require_auth)):
    """Extract structured data from multiple papers and return a comparison table."""
    try:
        extracted = extract_batch(request.papers, use_llm=request.use_llm)
        table = build_comparison_table(extracted)
        return {
            "extracted_data": [e.model_dump() for e in extracted],
            "comparison_table": table,
            "total_papers": len(extracted),
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Batch extraction error: {str(e)}")


# ---------------------------------------------------------------------------
# Drafting Assistant (requires auth)
# ---------------------------------------------------------------------------

@app.post("/api/draft", response_model=DraftResponse)
async def draft(request: DraftRequest, user=Depends(require_auth)):
    """Generate an academic draft section with citations and disclaimer."""
    try:
        record_stat("draft", {"section_type": request.section_type})
        # Template-based generation (always works, provides citations)
        draft_response = generate_draft(request)

        # Try AI-powered enhancement
        context_parts = []
        for i, p in enumerate(request.papers[:10], 1):
            if isinstance(p.authors, str):
                auth = p.authors
            elif p.authors:
                auth = ", ".join(a.name if hasattr(a, 'name') else str(a) for a in p.authors[:3])
            else:
                auth = "Unknown"
            context_parts.append(f"[{i}] {auth} ({p.year}). {p.title}. {p.journal or ''}.")
            if p.abstract:
                context_parts.append(f"    Abstract: {p.abstract[:300]}")
        context = "\n".join(context_parts)

        ai_text, ai_model = await generate_academic_text(
            topic=request.topic,
            section_type=request.section_type,
            context=context,
            max_words=request.max_words,
        )
        if ai_text:
            draft_response.content = ai_text
            draft_response.word_count = len(ai_text.split())
            draft_response.disclaimer = (
                f"AI-assisted content generated by {ai_model}. "
                "All AI-generated content must be reviewed, verified, and disclosed per journal policy."
            )
        return draft_response
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Drafting error: {str(e)}")


# ---------------------------------------------------------------------------
# Quick Generate: single prompt → final output (requires auth)
# ---------------------------------------------------------------------------

def _clean_search_query(query: str) -> str:
    """
    Clean and optimize a search query for academic APIs.
    Removes filler words, instruction phrases, and limits length.
    """
    q = query.strip()

    # Remove common instruction/filler phrases (case-insensitive)
    filler_phrases = [
        r"i\s+need\s+you\s+to\s+write\s+(?:a|an)?\s*",
        r"please\s+write\s+(?:a|an)?\s*",
        r"please\s+generate\s+(?:a|an)?\s*",
        r"write\s+(?:a|an)?\s*",
        r"generate\s+(?:a|an)?\s*",
        r"create\s+(?:a|an)?\s*",
        r"produce\s+(?:a|an)?\s*",
        r"draft\s+(?:a|an)?\s*",
        r"well\s+journal\s+high\s+quality\s*",
        r"well\s+journal\s*",
        r"high\s+quality\s*",
        r"journal\s+quality\s*",
        r"academic\s+quality\s*",
        r"q[1-4]\s+quality\s+style\s*",
        r"q[1-4]\s+quality\s*",
        r"q[1-4]\s+style\s*",
        r"with\s+this\s+topic\s*:\s*",
        r"topic\s*:\s*",
        r"with\s+q[1-4]\s*",
        r"in\s+\d+\s+words\s*",
        r"max\s+\d+\s+words\s*",
    ]
    for phrase in filler_phrases:
        q = re.sub(phrase, " ", q, flags=re.IGNORECASE)

    # Remove section type keywords
    for section_word in ["literature review", "lit review", "introduction", "abstract",
                          "conclusion", "summary", "section"]:
        q = re.sub(re.escape(section_word), " ", q, flags=re.IGNORECASE)

    # Remove common filler words (conservative — don't remove words that
    # could be part of meaningful academic terms like "institutional quality")
    filler_words = {
        "i", "need", "you", "to", "a", "an", "the", "please", "me", "my",
        "for", "about", "on", "with", "this", "that", "as", "is",
        "well", "topic", "type",
        "words", "max", "minimum", "maximum", "write", "generate", "create",
        "produce", "draft",
    }
    words = q.split()
    words = [w.strip(":,.;-") for w in words]
    meaningful = [w for w in words if w.lower() not in filler_words and len(w) > 1]
    cleaned = " ".join(meaningful)

    # If we stripped too much, fall back to original minus obvious phrases
    if len(cleaned) < 5:
        cleaned = re.sub(r"\s+", " ", q).strip(" :,-")

    # Limit to reasonable length for API queries
    if len(cleaned) > 120:
        cleaned = cleaned[:120].rsplit(" ", 1)[0]

    return cleaned.strip(" :,-")


def _parse_prompt(prompt: str) -> tuple[str, str]:
    """
    Parse a natural language prompt to extract section type and topic.
    Handles complex prompts with instructions, quality descriptors,
    and topic markers like "Topic: ..." or "about/on ...".

    Examples:
      "Write a literature review on FDI and growth"
        -> ("literature_review", "FDI and growth")
      "I need you to write a high quality Literature Review with Q1 style with this Topic: Food Insecurity in Liberia"
        -> ("literature_review", "Food Insecurity in Liberia")
      "Decades of Food Insecurity in Liberia: Staple Price Volatility"
        -> ("literature_review", "Decades of Food Insecurity in Liberia: Staple Price Volatility")
    """
    prompt_lower = prompt.lower().strip()

    section_map = {
        "literature review": "literature_review",
        "lit review": "literature_review",
        "introduction": "introduction",
        "abstract": "abstract",
        "conclusion": "conclusion",
        "summary": "summary",
    }

    # Detect section type
    section_type = "literature_review"  # default
    for key, val in section_map.items():
        if key in prompt_lower:
            section_type = val
            break

    # Strategy 1: Look for "Topic:" marker — very common in user prompts
    topic_match = re.search(r"topic\s*:\s*(.+)", prompt, re.IGNORECASE)
    if topic_match:
        raw_topic = topic_match.group(1).strip()
        cleaned = _clean_search_query(raw_topic)
        if cleaned and len(cleaned) > 3:
            return section_type, cleaned

    # Strategy 2: Look for "about/on/regarding [topic]" pattern
    about_match = re.search(
        r"(?:about|on|regarding|entitled)\s+[:\"]?(.+?)(?:\s+with\s+q\d|\s+in\s+\d+\s+words|$)",
        prompt, re.IGNORECASE,
    )
    if about_match:
        raw_topic = about_match.group(1).strip()
        cleaned = _clean_search_query(raw_topic)
        if cleaned and len(cleaned) > 3:
            return section_type, cleaned

    # Strategy 3: Remove section keywords and instruction phrases, use remainder
    topic = prompt
    for key in section_map:
        topic = re.sub(re.escape(key), " ", topic, flags=re.IGNORECASE)
    cleaned = _clean_search_query(topic)
    if cleaned and len(cleaned) > 5:
        return section_type, cleaned

    # Strategy 4: Default — clean the entire prompt
    return section_type, _clean_search_query(prompt)


@app.post("/api/research/quick-generate")
async def quick_generate(request: QuickGenerateRequest, user=Depends(require_auth)):
    """
    Single-prompt research generation.

    Takes a natural language prompt, automatically:
      1. Parses the topic and section type
      2. Searches for relevant papers
      3. Extracts structured data
      4. Generates a draft with citations
      5. Returns everything in one response
    """
    try:
        # Step 1: Use section_type from the dropdown (request), parse topic from prompt
        section_type = request.section_type or "literature_review"
        _, topic = _parse_prompt(request.prompt)
        record_stat("quick_generate", {"section_type": section_type})

        # Step 2: Search for papers (with fallback strategy)
        search_request = SearchRequest(
            query=topic,
            databases=request.databases,
            quartiles=request.quartiles,
            year_from=request.year_from,
            year_to=request.year_to,
            max_results=request.max_results,
        )
        search_response = await unified_search(search_request)
        papers = search_response.papers

        # Fallback 1: Try with a shorter query (first 5-6 key words)
        if not papers:
            words = topic.split()
            if len(words) > 4:
                short_topic = " ".join(words[:5])
                print(f"[QuickGenerate] Fallback 1: trying shorter query '{short_topic}'")
                fallback_request = SearchRequest(
                    query=short_topic,
                    databases=request.databases,
                    quartiles=request.quartiles,
                    year_from=request.year_from,
                    year_to=request.year_to,
                    max_results=request.max_results,
                )
                search_response = await unified_search(fallback_request)
                papers = search_response.papers
                if papers:
                    topic = short_topic

        # Fallback 2: Try with even shorter query (first 3 words)
        if not papers:
            words = topic.split()
            if len(words) > 2:
                short_topic = " ".join(words[:3])
                print(f"[QuickGenerate] Fallback 2: trying shorter query '{short_topic}'")
                fallback_request = SearchRequest(
                    query=short_topic,
                    databases=request.databases,
                    quartiles=request.quartiles,
                    year_from=request.year_from,
                    year_to=request.year_to,
                    max_results=request.max_results,
                )
                search_response = await unified_search(fallback_request)
                papers = search_response.papers
                if papers:
                    topic = short_topic

        # Fallback 3: Try without quartile filter (accept all quartiles)
        if not papers:
            print(f"[QuickGenerate] Fallback 3: trying without quartile filter")
            broader_request = SearchRequest(
                query=topic,
                databases=request.databases,
                quartiles=[],  # Empty = no quartile filtering
                year_from=request.year_from,
                year_to=request.year_to,
                max_results=request.max_results,
            )
            search_response = await unified_search(broader_request)
            papers = search_response.papers

        # Fallback 4: Try with a broader topic (first 2 words only)
        if not papers:
            words = topic.split()
            if len(words) >= 2:
                broad_topic = " ".join(words[:2])
                print(f"[QuickGenerate] Fallback 4: trying broad query '{broad_topic}'")
                broad_request = SearchRequest(
                    query=broad_topic,
                    databases=request.databases,
                    quartiles=[],
                    year_from=request.year_from,
                    year_to=request.year_to,
                    max_results=request.max_results,
                )
                search_response = await unified_search(broad_request)
                papers = search_response.papers
                if papers:
                    topic = broad_topic

        if not papers:
            return QuickGenerateResponse(
                topic=topic,
                section_type=section_type,
                papers=[],
                extracted_data=[],
                comparison_table=[],
                draft=DraftResponse(
                    section_type=section_type,
                    topic=topic,
                    content="No papers were found for this topic. Please try a different search query.",
                    citations=[],
                    word_count=0,
                    disclaimer="No content generated — no sources found.",
                ),
                disclosure="No AI-assisted content was generated.",
                total_sources=0,
            )

        # Step 3: Extract data from papers
        extracted = extract_batch(papers, use_llm=False)
        comparison_table = build_comparison_table(extracted)

        # Step 4: Generate draft — template first (for citations), then AI enhancement
        draft_request = DraftRequest(
            section_type=section_type,
            topic=topic,
            papers=papers,
            extracted_data=extracted,
            style="academic",
            use_llm=False,
            max_words=request.max_words,
        )
        draft_response = generate_draft(draft_request)

        # Try AI-powered generation (auto-routes to best available model)
        context_parts = []
        for i, p in enumerate(papers[:10], 1):
            if isinstance(p.authors, str):
                auth = p.authors
            elif p.authors:
                auth = ", ".join(a.name if hasattr(a, 'name') else str(a) for a in p.authors[:3])
            else:
                auth = "Unknown"
            context_parts.append(f"[{i}] {auth} ({p.year}). {p.title}. {p.journal or ''}.")
            if p.abstract:
                context_parts.append(f"    Abstract: {p.abstract[:300]}")
        context = "\n".join(context_parts)

        ai_text, ai_model = await generate_academic_text(
            topic=topic,
            section_type=section_type,
            context=context,
            max_words=request.max_words,
        )
        if ai_text:
            draft_response.content = ai_text
            draft_response.word_count = len(ai_text.split())
            draft_response.disclaimer = (
                f"AI-assisted content generated by {ai_model}. "
                "All AI-generated content must be reviewed, verified, and disclosed per journal policy."
            )

        # Step 5: Return consolidated result
        return QuickGenerateResponse(
            topic=topic,
            section_type=section_type,
            papers=papers,
            extracted_data=extracted,
            comparison_table=comparison_table,
            draft=draft_response,
            disclosure=draft_response.disclaimer,
            total_sources=len(papers),
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Quick-generate error: {str(e)}")


# ---------------------------------------------------------------------------
# Provenance & Integrity Audit (requires auth)
# ---------------------------------------------------------------------------

@app.post("/api/audit/session")
async def create_audit_session(user=Depends(require_auth)):
    session = create_session()
    return {"session_id": session.session_id, "created_at": session.created_at}


@app.get("/api/audit/sessions")
async def list_audit_sessions(user=Depends(require_auth)):
    return {"sessions": list_sessions()}


@app.get("/api/audit/{session_id}")
async def get_audit_report(session_id: str, user=Depends(require_auth)):
    session = get_session(session_id)
    if not session:
        raise HTTPException(status_code=404, detail=f"Session '{session_id}' not found")
    return session.to_dict()


@app.post("/api/audit/{session_id}/verify")
async def mark_verified(session_id: str, all_records: bool = True, index: int = 0,
                        user=Depends(require_auth)):
    session = get_session(session_id)
    if not session:
        raise HTTPException(status_code=404, detail=f"Session '{session_id}' not found")
    if all_records:
        session.mark_all_verified()
    else:
        session.mark_verified(index)
    return {"status": "verified", "session_id": session_id}


# ---------------------------------------------------------------------------
# User Profile Management (requires auth)
# ---------------------------------------------------------------------------

@app.put("/api/auth/profile")
async def update_profile(request: ProfileUpdateRequest, user=Depends(require_auth)):
    """Update the current user's profile."""
    try:
        updated = update_user_profile(
            user_id=user.id,
            name=request.name,
            institution=request.institution,
            research_field=request.research_field,
        )
        return updated
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.post("/api/auth/change-password")
async def change_user_password(request: PasswordChangeRequest, user=Depends(require_auth)):
    """Change the current user's password."""
    try:
        change_password(user.id, request.old_password, request.new_password)
        return {"status": "password_changed"}
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


# ---------------------------------------------------------------------------
# Admin Routes (requires admin role)
# ---------------------------------------------------------------------------

@app.get("/api/admin/users")
async def admin_list_users(admin=Depends(require_admin)):
    """List all registered users (admin only)."""
    return {"users": get_all_users()}


@app.get("/api/admin/users/{user_id}")
async def admin_get_user(user_id: str, admin=Depends(require_admin)):
    """Get a specific user's details (admin only)."""
    user = get_user_by_id(user_id)
    if not user:
        raise HTTPException(status_code=404, detail="User not found.")
    return user


@app.put("/api/admin/users/{user_id}/role")
async def admin_update_user_role(user_id: str, request: UserRoleUpdateRequest,
                                  admin=Depends(require_admin)):
    """Update a user's role (admin only)."""
    try:
        return update_user_role(user_id, request.role)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.post("/api/admin/users/{user_id}/toggle-active")
async def admin_toggle_user_active(user_id: str, admin=Depends(require_admin)):
    """Activate or deactivate a user account (admin only)."""
    try:
        return toggle_user_active(user_id)
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))


@app.delete("/api/admin/users/{user_id}")
async def admin_delete_user(user_id: str, admin=Depends(require_admin)):
    """Delete a user account (admin only)."""
    try:
        delete_user(user_id)
        return {"status": "deleted", "user_id": user_id}
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))


@app.post("/api/admin/users/{user_id}/reset-password")
async def admin_reset_user_password(request: AdminResetPasswordRequest,
                                     admin=Depends(require_admin)):
    """Reset any user's password (admin only)."""
    try:
        admin_reset_password(request.user_id, request.new_password)
        return {"status": "password_reset", "user_id": request.user_id}
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.get("/api/admin/config")
async def admin_get_config(admin=Depends(require_admin)):
    """Get system configuration (admin only)."""
    return get_config()


@app.put("/api/admin/config")
async def admin_update_config(request: ConfigUpdateRequest, admin=Depends(require_admin)):
    """Update system configuration (admin only)."""
    return update_config(request.updates)


@app.get("/api/admin/ai-config")
async def admin_get_ai_config(admin=Depends(require_admin)):
    """Get AI model provider status, API key configuration, and task routing (admin only).

    Returns each provider's display name, available models, free/paid tier,
    whether a key is configured, the key source (admin config vs env var),
    and a masked preview of the key. Also returns the full task routing map.
    """
    return get_ai_status()


@app.post("/api/admin/ai-config/test")
async def admin_test_ai_provider(request: AiTestRequest, admin=Depends(require_admin)):
    """Test an AI provider's API key by making a minimal request (admin only).

    If `api_key` is provided, tests that key directly without saving it.
    Otherwise, tests the currently configured key for the provider.
    """
    result = await test_provider_connection(request.provider, request.api_key)
    return result


@app.get("/api/admin/feature-flags")
async def admin_get_feature_flags(admin=Depends(require_admin)):
    """Get feature toggle settings (admin only)."""
    return get_feature_flags()


@app.get("/api/admin/content")
async def admin_get_content(key: str = None, admin=Depends(require_admin)):
    """Get editable content (admin only)."""
    return get_content(key)


@app.put("/api/admin/content")
async def admin_update_content(request: ContentUpdateRequest, admin=Depends(require_admin)):
    """Update editable content (admin only)."""
    return update_content(request.key, request.content)


@app.get("/api/admin/content/keys")
async def admin_list_content_keys(admin=Depends(require_admin)):
    """List all editable content keys (admin only)."""
    return {"keys": list_content_keys()}


@app.get("/api/admin/stats")
async def admin_get_stats(admin=Depends(require_admin)):
    """Get platform statistics (admin only)."""
    return get_stats()


@app.post("/api/admin/stats/reset")
async def admin_reset_stats(admin=Depends(require_admin)):
    """Reset platform statistics (admin only)."""
    return reset_stats()


@app.get("/api/admin/database-sources")
async def admin_get_database_sources(admin=Depends(require_admin)):
    """Get all database sources (admin only)."""
    return {"sources": get_database_sources()}


@app.post("/api/admin/database-sources/{source_id}/toggle")
async def admin_toggle_database_source(source_id: str, admin=Depends(require_admin)):
    """Enable or disable a database source (admin only)."""
    try:
        return toggle_database_source(source_id)
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))


@app.get("/api/admin/tools")
async def admin_get_tools(category: str = None, admin=Depends(require_admin)):
    """Get all tool integrations (admin only)."""
    return {"tools": get_tool_integrations(category)}


@app.post("/api/admin/tools/{tool_id}/toggle")
async def admin_toggle_tool(tool_id: str, admin=Depends(require_admin)):
    """Enable or disable a tool integration (admin only)."""
    try:
        return toggle_tool_integration(tool_id)
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))


# ---------------------------------------------------------------------------
# Data Analysis: R / RStudio / Stata 18 (requires auth)
# ---------------------------------------------------------------------------

@app.get("/api/analysis/methods")
async def analysis_methods(user=Depends(require_auth)):
    """Get all available analysis methods."""
    return {"methods": get_available_methods()}


@app.post("/api/analysis/generate", response_model=AnalysisResponse)
async def analysis_generate_code(request: AnalysisRequest, user=Depends(require_auth)):
    """Generate R or Stata 18 code for the specified analysis method."""
    variables = {
        "dependent": request.dependent_variable,
        "independent": request.independent_variables,
        "controls": request.control_variables,
        "data_file": request.data_file,
        "entity": request.entity_variable,
        "time": request.time_variable,
        "instruments": request.instruments,
        "treatment": request.treatment_variable,
        "post": request.post_variable,
    }
    options = {
        "robust_se": request.robust_se,
        "cluster_var": request.cluster_variable,
        **request.options,
    }

    if request.language.lower() == "stata":
        code = generate_stata_code(request.method, variables, options)
        description = ANALYSIS_METHODS.get(request.method, {}).get("description", "")
        packages = ["estout", "reghdfe", "coefplot"]
    else:
        code = generate_r_code(request.method, variables, options)
        description = ANALYSIS_METHODS.get(request.method, {}).get("description", "")
        packages = ["tidyverse", "sandwich", "lmtest", "stargazer"]

    return AnalysisResponse(
        method=request.method,
        language=request.language,
        code=code,
        description=description,
        packages_required=packages,
    )


@app.post("/api/analysis/recommend")
async def analysis_recommend(request: AnalysisRecommendationRequest, user=Depends(require_auth)):
    """Get analysis method recommendations based on research context."""
    recommendations = recommend_analysis(
        research_topic=request.research_topic,
        variables=request.variables,
        data_type=request.data_type,
    )
    return {"recommendations": recommendations}


# ---------------------------------------------------------------------------
# Writing Tools: Grammarly, QuillBot, Paperpal (requires auth)
# ---------------------------------------------------------------------------

@app.get("/api/writing/tools")
async def writing_tools_list(user=Depends(require_auth)):
    """Get all writing tool integrations."""
    return {"tools": get_writing_tools()}


@app.get("/api/writing/tools/{tool_id}")
async def writing_tool_detail(tool_id: str, user=Depends(require_auth)):
    """Get details of a specific writing tool."""
    tool = get_writing_tool(tool_id)
    if not tool:
        raise HTTPException(status_code=404, detail=f"Writing tool '{tool_id}' not found.")
    return tool


@app.post("/api/writing/analyze")
async def writing_analyze(request: WritingAnalysisRequest, user=Depends(require_auth)):
    """Analyze writing quality and provide improvement suggestions."""
    return analyze_writing(request.text)


@app.post("/api/writing/journal-check")
async def writing_journal_check(request: WritingAnalysisRequest, user=Depends(require_auth)):
    """Check writing readiness for journal submission."""
    return enhance_for_journal(request.text, request.journal_name or None)


@app.post("/api/writing/fix-grammar")
async def writing_fix_grammar(request: WritingAnalysisRequest, user=Depends(require_auth)):
    """Fix grammar and spelling issues directly on the platform (Grammarly-style)."""
    # Try AI-powered grammar fix first
    ai_text, ai_model = await enhance_text_with_ai(request.text, "grammar_fix")
    if ai_text:
        return {
            "original": request.text,
            "enhanced": ai_text,
            "changes": [{"type": "ai_enhanced", "description": f"Grammar fixed by {ai_model}"}],
            "summary": f"Grammar and spelling corrected by {ai_model}.",
            "model_used": ai_model,
        }
    # Fall back to local regex-based engine
    return fix_grammar(request.text)


@app.post("/api/writing/paraphrase")
async def writing_paraphrase(request: WritingAnalysisRequest, user=Depends(require_auth)):
    """Paraphrase text for clarity and conciseness (QuillBot-style)."""
    # Try AI-powered paraphrasing first
    ai_text, ai_model = await enhance_text_with_ai(request.text, "paraphrase")
    if ai_text:
        return {
            "original": request.text,
            "enhanced": ai_text,
            "changes": [{"type": "ai_enhanced", "description": f"Paraphrased by {ai_model}"}],
            "summary": f"Text paraphrased by {ai_model}.",
            "model_used": ai_model,
        }
    # Fall back to local engine
    return paraphrase_text(request.text)


@app.post("/api/writing/enhance-academic")
async def writing_enhance_academic(request: WritingAnalysisRequest, user=Depends(require_auth)):
    """Enhance academic tone and vocabulary (Paperpal-style)."""
    # Try AI-powered academic enhancement first
    ai_text, ai_model = await enhance_text_with_ai(request.text, "academic_enhance")
    if ai_text:
        return {
            "original": request.text,
            "enhanced": ai_text,
            "changes": [{"type": "ai_enhanced", "description": f"Academically enhanced by {ai_model}"}],
            "summary": f"Academic tone enhanced by {ai_model}.",
            "model_used": ai_model,
        }
    # Fall back to local engine
    return enhance_academic(request.text)


@app.post("/api/writing/enhance-all")
async def writing_enhance_all(request: WritingAnalysisRequest, user=Depends(require_auth)):
    """Apply all three enhancement tools in sequence: grammar, paraphrase, academic."""
    # Try AI-powered enhancement (academic_enhance covers all three aspects)
    ai_text, ai_model = await enhance_text_with_ai(request.text, "academic_enhance")
    if ai_text:
        return {
            "original": request.text,
            "enhanced": ai_text,
            "changes": [{"type": "ai_enhanced", "description": f"Full enhancement by {ai_model}"}],
            "summary": f"Grammar, paraphrasing, and academic enhancement applied by {ai_model}.",
            "model_used": ai_model,
        }
    # Fall back to local engine
    return enhance_text_all(request.text)


@app.post("/api/writing/humanize")
async def writing_humanize(request: WritingAnalysisRequest, user=Depends(require_auth)):
    """Humanize and naturalize text by reducing formulaic AI patterns.

    Improves readability and authentic voice by:
    - Reducing overused AI transition words
    - Varying sentence beginnings
    - Replacing formulaic phrases with natural alternatives
    - Varying sentence length for natural rhythm

    AI provenance is still tracked and disclosed per platform integrity policy.
    """
    # Try AI-powered humanization first
    ai_text, ai_model = await enhance_text_with_ai(request.text, "paraphrase")
    if ai_text:
        # Run local humanizer on AI output for best results
        result = humanize_text(ai_text)
        result["model_used"] = ai_model
        result["summary"] = (
            f"Text humanized and naturalized (AI-assisted by {ai_model}, "
            f"then locally refined). {len(result.get('changes', []))} "
            f"pattern(s) reduced. Naturalness score: {result.get('naturalness_score', 0)}/100."
        )
        return result
    # Fall back to local engine only
    return humanize_text(request.text)


# ---------------------------------------------------------------------------
# Reference Management: Mendeley, Zotero, NotePal (requires auth)
# ---------------------------------------------------------------------------

@app.get("/api/references/managers")
async def reference_managers_list(user=Depends(require_auth)):
    """Get all reference manager integrations."""
    return {"managers": get_reference_managers()}


@app.get("/api/references/managers/{manager_id}")
async def reference_manager_detail(manager_id: str, user=Depends(require_auth)):
    """Get details of a specific reference manager."""
    manager = get_reference_manager(manager_id)
    if not manager:
        raise HTTPException(status_code=404, detail=f"Reference manager '{manager_id}' not found.")
    return manager


@app.get("/api/references/formats")
async def reference_formats_list(user=Depends(require_auth)):
    """Get all available export formats."""
    return {"formats": get_export_formats()}


@app.post("/api/references/export", response_model=ReferenceExportResponse)
async def references_export(request: ReferenceExportRequest, user=Depends(require_auth)):
    """Export papers in the specified reference format."""
    if not request.papers:
        raise HTTPException(status_code=400, detail="No papers provided for export.")
    return export_references(request.papers, request.format)


@app.post("/api/references/cite")
async def references_cite(request: CitationFormatRequest, user=Depends(require_auth)):
    """Format papers as citations in the specified style."""
    if not request.papers:
        raise HTTPException(status_code=400, detail="No papers provided.")
    citations = format_citations(request.papers, request.style)
    return {"style": request.style, "citations": citations, "count": len(citations)}


@app.post("/api/references/verify-citations")
async def references_verify_citations(request: CitationVerificationRequest,
                                       user=Depends(require_auth)):
    """Verify that inline citations in the text match the reference list.

    Flags orphan citations (in text but not in references), uncited references
    (in reference list but never cited), and mismatches (wrong author/year).
    """
    result = verify_citations(request.text, request.references)
    return result


# ---------------------------------------------------------------------------
# Public Info (no auth required)
# ---------------------------------------------------------------------------

@app.get("/api/features")
async def features():
    """Get platform features and capabilities (public)."""
    return {
        "platform": "THEeye",
        "version": "3.0.0",
        "features": get_feature_flags(),
        "databases": get_database_sources(),
        "tools": get_tool_integrations(),
        "analysis_methods": get_available_methods(),
        "writing_tools": get_writing_tools(),
        "reference_managers": get_reference_managers(),
        "export_formats": get_export_formats(),
    }


# ---------------------------------------------------------------------------
# Document Export
# ---------------------------------------------------------------------------

@app.post("/api/research/export-document")
async def export_document_endpoint(request: DocumentExportRequest, user=Depends(require_auth)):
    """Export generated research text to Word, PDF, HTML, Markdown, or TXT."""
    from fastapi.responses import Response as _Response
    try:
        file_bytes, filename, mime_type = _export_document(
            text=request.text,
            fmt=request.format,
            title=request.title or "THEeye Research Output",
            citations=request.citations,
            disclaimer=request.disclaimer,
            section_type=request.section_type,
            topic=request.topic,
            word_count=request.word_count,
            total_sources=request.total_sources,
        )
        headers = {"Content-Disposition": 'attachment; filename="{}"'.format(filename)}
        return _Response(content=file_bytes, media_type=mime_type, headers=headers)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail="Export failed: {}".format(str(e)))


# ---------------------------------------------------------------------------
# Document Drafting & Roadmap Engine (requires auth)
# ---------------------------------------------------------------------------

@app.get("/api/drafting/document-types")
async def drafting_document_types(user=Depends(require_auth)):
    """Get all available document types and their format options."""
    return {"document_types": get_document_types()}


@app.get("/api/drafting/research-fields")
async def drafting_research_fields(user=Depends(require_auth)):
    """Get all available research fields for topic suggestion."""
    return {"research_fields": get_research_fields()}


@app.post("/api/drafting/roadmap")
async def drafting_roadmap(request: RoadmapRequest, user=Depends(require_auth)):
    """Generate a full document roadmap (structure, sections, word estimates, writing guidelines)
    for research articles, theses, literature reviews, book reports, review papers, etc.
    Supports Q1-Q4 journals and thesis formats from Chinese, USA, European, African, and UK traditions."""
    try:
        record_stat("drafting_roadmap", {"document_type": request.document_type})
        # Template-based generation (always works)
        result = generate_roadmap(
            document_type=request.document_type,
            format=request.format,
            topic=request.topic,
            field=request.field,
        )
        if "error" in result:
            raise HTTPException(status_code=400, detail=result["error"])

        # Try AI-powered roadmap generation
        ai_prompt = (
            f"Create a detailed document roadmap for a {request.document_type.replace('_', ' ')} "
            f"in {request.format} format on the topic: '{request.topic}'."
            f"{f' Research field: {request.field}.' if request.field else ''}\n\n"
            f"Include for each section: title, purpose, estimated word count, and writing guidelines.\n"
            f"Return the roadmap in Markdown format with clear section headings."
        )
        ai_text, ai_model = await generate_text(
            user_prompt=ai_prompt,
            task_type="roadmap_generation",
            max_tokens=2500,
            temperature=0.5,
        )
        if ai_text:
            result["ai_roadmap"] = ai_text
            result["model_used"] = ai_model

        return result
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Roadmap generation error: {str(e)}")


@app.post("/api/drafting/topics")
async def drafting_topics(request: TopicSuggestionRequest, user=Depends(require_auth)):
    """Suggest novel research topics from all fields of study.
    Prioritizes topics that have not been researched before, with novelty analysis,
    methodology recommendations, and target journals."""
    try:
        record_stat("drafting_topics", {"field": request.field})
        # Template-based generation (always works)
        result = suggest_research_topics(
            field=request.field,
            keywords=request.keywords,
            focus_novelty=request.focus_novelty,
            max_topics=request.max_topics,
        )

        # Try AI-powered topic suggestion
        ai_prompt = (
            f"Suggest {request.max_topics} novel research topics in the field of {request.field}."
            f"{f' Related keywords: {request.keywords}.' if request.keywords else ''}\n\n"
            f"For each topic, provide:\n"
            f"1. The topic title\n"
            f"2. Why it matters (research gap it addresses)\n"
            f"3. Suggested methodology\n"
            f"4. A novelty score (1-10) with brief justification\n\n"
            f"Focus on topics that are genuinely under-researched and feasible for a graduate researcher."
        )
        ai_text, ai_model = await generate_text(
            user_prompt=ai_prompt,
            task_type="topic_suggestion",
            max_tokens=2500,
            temperature=0.8,
        )
        if ai_text:
            result["ai_topics"] = ai_text
            result["model_used"] = ai_model

        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Topic suggestion error: {str(e)}")


@app.post("/api/drafting/open-access-articles")
async def drafting_open_access_articles(request: OpenAccessArticleRequest, user=Depends(require_auth)):
    """Find open-access articles from all OA journal platforms that best suit a given topic.
    Searches OpenAlex and Semantic Scholar for papers with downloadable PDFs."""
    try:
        record_stat("drafting_oa_articles", {"topic": request.topic[:50]})
        result = await find_open_access_articles(
            topic=request.topic,
            max_results=request.max_results,
        )
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Open-access article search error: {str(e)}")


@app.post("/api/drafting/download-pdf")
async def drafting_download_pdf(request: DownloadPdfRequest, user=Depends(require_auth)):
    """Download a PDF from an open-access article URL directly through THEeye.
    Returns the PDF file as a binary download."""
    try:
        content, filename, content_type = await download_pdf(request.url)
        headers = {"Content-Disposition": f'attachment; filename="{filename}"'}
        return Response(content=content, media_type=content_type, headers=headers)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF download failed: {str(e)}")


@app.get("/api/drafting/oa-sources")
async def drafting_oa_sources(user=Depends(require_auth)):
    """Return the list of open-access paper sources with paper counts."""
    return {"sources": OA_SOURCES, "total_papers": "300M+"}


# ---------------------------------------------------------------------------
# AI Router Info (requires auth)
# ---------------------------------------------------------------------------

@app.get("/api/ai/models")
async def ai_models_info(user=Depends(require_auth)):
    """Return available AI models, providers, and task routing configuration."""
    return {
        "available_providers": get_available_providers(),
        "routing": get_routing_info(),
    }


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("backend.main:app", host="0.0.0.0", port=80, reload=True)

# ---------------------------------------------------------------------------
# Data Analysis Endpoints (Advanced)
# ---------------------------------------------------------------------------
@app.post("/api/analysis/upload-dataset")
async def api_upload_dataset(request: Request):
    auth_check = _require_auth(request)
    if auth_check: return auth_check
    form = await request.form()
    file = form.get("file")
    if not file: raise HTTPException(status_code=400, detail="No file uploaded")
    file_bytes = await file.read()
    try:
        result = parse_dataset(file_bytes, file.filename)
        return {"status": "success", "dataset": result}
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))
@app.post("/api/analysis/run")
async def api_run_analysis(request: Request):
    auth_check = _require_auth(request)
    if auth_check: return auth_check
    body = await request.json()
    instruction = body.get("instruction", "")
    tool = body.get("tool", "python")
    dataset_path = body.get("dataset_path", os.path.join(tempfile.gettempdir(), "theeye_current_dataset.csv"))
    if not os.path.exists(dataset_path):
        raise HTTPException(status_code=400, detail="No dataset uploaded.")
    try:
        result = execute_analysis(instruction, dataset_path, tool)
        return {"status": "success", "result": result}
    except Exception as e:
        import traceback; traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))
@app.post("/api/analysis/online-data")
async def api_online_data(request: Request):
    auth_check = _require_auth(request)
    if auth_check: return auth_check
    body = await request.json()
    try:
        result = extract_online_data(body.get("source", "world_bank"), body.get("query", ""), body.get("params", {}))
        return {"status": "success", "data": result}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
@app.post("/api/analysis/suggest-model")
async def api_suggest_model(request: Request):
    auth_check = _require_auth(request)
    if auth_check: return auth_check
    body = await request.json()
    methodology = body.get("methodology", "")
    if not methodology.strip():
        raise HTTPException(status_code=400, detail="Please paste your methodology.")
    try:
        result = suggest_econometric_model(methodology)
        return {"status": "success", "suggestion": result}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
@app.post("/api/analysis/generate-code")
async def api_generate_code(request: Request):
    auth_check = _require_auth(request)
    if auth_check: return auth_check
    body = await request.json()
    tool = body.get("tool", "python")
    method = body.get("method", "ols")
    variables = body.get("variables", {})
    try:
        if tool in ("r", "rstudio"): code = generate_r_code(method, variables)
        elif tool == "stata": code = generate_stata_code(method, variables)
        else: code = generate_tool_code(tool, method, variables.get("independent", []), [])
        return {"status": "success", "code": code}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
